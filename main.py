import os
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from PIL import Image
import pytesseract
from googletrans import Translator
import fitz
import tempfile
import cv2
import numpy as np
import json
import webbrowser
from PIL import Image

# Load GUI text from JSON file


def load_gui_text(json_file_path):
    gui_text = {}
    if os.path.exists(json_file_path):
        try:
            with open(json_file_path, "r", encoding="utf-8") as json_file:
                gui_text = json.load(json_file)
        except Exception as e:
            print(f"Error loading JSON file: {e}")
    else:
        print(f"JSON file not found: {json_file_path}")
    return gui_text


# Function to extract images from PDF using PyMuPDF
def extract_images_from_pdf(pdf_path, image_folder):
    images = []
    try:
        with fitz.open(pdf_path) as pdf_document:
            for page_number, page in enumerate(pdf_document):
                image_list = page.get_images(full=True)
                if image_list:
                    for image_index, img in enumerate(page.get_images(full=True)):
                        xref = img[0]
                        base_image = pdf_document.extract_image(xref)
                        if "image" in base_image:  # Check if image data is available
                            image_bytes = base_image["image"]
                            image = cv2.imdecode(np.frombuffer(
                                image_bytes, np.uint8), cv2.IMREAD_COLOR)
                            if image is not None:
                                image_path = os.path.join(image_folder, f"page_{
                                                          page_number + 1}_image_{image_index + 1}.png")
                                cv2.imwrite(image_path, image)
                                images.append(image_path)
    except Exception as e:
        print(f"Error extracting images from PDF: {e}")
    return images


# Function to convert PDF to PNG images
def pdf_to_png(pdf_path):
    png_paths = []
    try:
        with fitz.open(pdf_path) as pdf_document:
            for page_number, page in enumerate(pdf_document):
                pix = page.get_pixmap()
                output_path = os.path.join(tempfile.gettempdir(), f"page_{
                                           page_number + 1}.png")
                # Use write_image method instead of writePNG
                pix.pil_save(output_path)
                png_paths.append(output_path)
    except Exception as e:
        print(f"Error converting PDF to PNG images: {e}")
    return png_paths


# Function to perform OCR on PNG images
def ocr_png(png_paths):
    ocr_text = ""
    for png_path in png_paths:
        text = pytesseract.image_to_string(Image.open(png_path))
        ocr_text += text + "\n"
    return ocr_text


# Function to translate text using Google Translate API
def translate_text(text, target_language='fa'):
    translator = Translator()
    translated_text = translator.translate(text, dest=target_language)
    return translated_text.text


# Function to save translated content to Word document
def save_to_word(text, image_paths, output_file):
    doc = Document()
    doc.add_paragraph(text)
    for image_path in image_paths:
        doc.add_picture(image_path)
    doc.save(output_file)


# Function to handle translation process triggered by GUI
def translate_from_gui():
    input_pdf = input_pdf_entry.get()
    output_word = output_word_entry.get()

    # Check if input and output paths are provided
    if not input_pdf or not output_word:
        messagebox.showerror(
            gui_text["en"]["messages"]["error"], gui_text["en"]["messages"]["error_no_paths"])
        return

    try:
        # Create a temporary folder to store intermediate files
        temp_folder = tempfile.mkdtemp()

        # Convert PDF to PNG images
        png_paths = pdf_to_png(input_pdf)

        # Perform OCR on PNG images
        ocr_text = ocr_png(png_paths)

        # Translate text
        translated_text = translate_text(ocr_text)

        # Extract images from PDF
        image_folder = os.path.join(temp_folder, "images")
        os.makedirs(image_folder, exist_ok=True)
        image_paths = extract_images_from_pdf(input_pdf, image_folder)

        # Save translated content to Word document
        save_to_word(translated_text, image_paths, output_word)

        # Show success message
        messagebox.showinfo(gui_text["en"]["messages"]["success"],
                            gui_text["en"]["messages"]["translation_completed"])

        # Open destination folder
        os.startfile(os.path.dirname(os.path.abspath(output_word)))

    except Exception as e:
        # Show error message if translation fails
        messagebox.showerror(gui_text["en"]["messages"]["error"], f"{
                             gui_text["en"]['messages']['error_occurred']} {str(e)}")


# Function to open file dialog for selecting PDF file
def choose_pdf_file():
    file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
    if file_path:
        input_pdf_entry.delete(0, tk.END)
        input_pdf_entry.insert(0, file_path)


# Function to open file dialog for selecting output Word file
def choose_word_file():
    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx", filetypes=[("Word files", "*.docx")])
    if file_path:
        output_word_entry.delete(0, tk.END)
        output_word_entry.insert(0, file_path)


# Function to open GitHub repository link
def open_github_repo():
    webbrowser.open(gui_text["en"]["github_repositroy"])


# Load GUI text from JSON file
gui_text = load_gui_text("./assets/gui.json")

# Create GUI window
root = tk.Tk()
root.title(gui_text["en"]["software_name"])

# Set program icon
root.iconbitmap("./assets/fav.ico")

# Input PDF file entry
input_pdf_label = tk.Label(root, text=gui_text["en"]["labels"]["input_pdf"])
input_pdf_label.grid(row=0, column=0, sticky="w")
input_pdf_entry = tk.Entry(root, width=50)
input_pdf_entry.grid(row=0, column=1, padx=5, pady=5)
choose_pdf_button = tk.Button(
    root, text=gui_text["en"]["buttons"]["choose_pdf"], command=choose_pdf_file)
choose_pdf_button.grid(row=0, column=2, padx=5, pady=5)

# Output Word file entry
output_word_label = tk.Label(
    root, text=gui_text["en"]["labels"]["output_word"])
output_word_label.grid(row=1, column=0, sticky="w")
output_word_entry = tk.Entry(root, width=50)
output_word_entry.grid(row=1, column=1, padx=5, pady=5)
choose_word_button = tk.Button(
    root, text=gui_text["en"]["buttons"]["choose_word"], command=choose_word_file)
choose_word_button.grid(row=1, column=2, padx=5, pady=5)

# Translate button
translate_button = tk.Button(
    root, text=gui_text["en"]["buttons"]["translate"], command=translate_from_gui, cursor="hand2")
translate_button.grid(row=2, column=0, columnspan=3, pady=10)

# Copyright notice
copyright_label = tk.Label(root, text=gui_text["en"]["copyright"])
copyright_label.grid(row=3, column=0, columnspan=2, pady=5)

# Create a label with a GitHub repository link
github_label = tk.Label(
    root, text=gui_text["en"]["labels"]["github_link"], fg="blue", cursor="hand2")
github_label.grid(row=3, column=2, pady=5)
github_label.bind("<Button-1>", lambda e: open_github_repo())

# Run the GUI
root.mainloop()
